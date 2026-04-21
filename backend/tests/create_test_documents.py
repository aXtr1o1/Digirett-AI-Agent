"""
create_test_documents.py

Generates 4 test documents in DOCX format for testing document upload pipeline.
Each document is a narrative-based company scenario related to lovdata.no domains.

Documents are stored in: backend/tests/test_docs/
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_document_1_tech_company():
    """
    Document 1: Tech Company Compliance Scenario
    Company: XYZ Software AS
    Topics: Data protection, employee management, compliance
    Related to: Employment & Data Protection law
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Company Profile & Compliance Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Organization Information', 1)
    doc.add_paragraph('Company Name: XYZ Software AS (Norwegian Tech Company)')
    doc.add_paragraph('Establishment Year: 2015')
    doc.add_paragraph('Number of Employees: 87')
    doc.add_paragraph('Industry: Software Development & Digital Solutions')
    doc.add_paragraph('Business Location: Oslo, Norway')
    
    doc.add_heading('Company Overview', 2)
    doc.add_paragraph(
        'XYZ Software AS is a mid-sized technology company specializing in cloud-based business solutions. '
        'The company develops and maintains custom software platforms for Norwegian businesses, focusing on data security '
        'and compliance with strict regulations. With 87 employees spread across development, sales, and operations teams, '
        'the company operates under rigorous data protection and employment practices.'
    )
    
    doc.add_heading('Employee Management Practices', 2)
    
    doc.add_heading('Employment Categories', 3)
    doc.add_paragraph(
        'Full-time Permanent Staff: 65 employees\n'
        '- Software developers and architects\n'
        '- Operations and infrastructure specialists\n'
        '- Sales and customer success representatives\n'
        '\n'
        'Contract-based Developers: 15 employees\n'
        '- Project-specific contractor arrangements\n'
        '- Average contract duration: 6-12 months\n'
        '\n'
        'Consultants and Advisors: 7 individuals\n'
        '- External expertise on regulatory matters\n'
        '- Part-time advisory capacity'
    )
    
    doc.add_heading('Working Hours & Schedule', 3)
    doc.add_paragraph(
        'Standard working hours: 37.5 hours per week (Monday-Friday, 09:00-17:00 with one-hour lunch break)\n'
        'Flexible arrangements: 22 employees on flexible schedules with core hours 10:00-15:00\n'
        'Remote work policy: 3 days office, 2 days remote per week (standard)\n'
        'Overtime handling: Compensated with time off (1.5x) or additional salary (1.1x base rate)\n'
        'Vacation: 25 days per year plus 5 public holidays'
    )
    
    doc.add_heading('Employee Rights & Responsibilities', 3)
    doc.add_paragraph(
        'Salary: Competitive market rates with annual review cycles\n'
        'Benefits: Health insurance, occupational pension (8%), professional development stipend\n'
        'Work Equipment: Company-provided laptops, dual monitors, office setup\n'
        'Confidentiality: All employees sign data protection and confidentiality agreements\n'
        'Performance Reviews: Semi-annual evaluations with goal-setting and feedback sessions\n'
        'Training & Development: Annual budget of 5,000 NOK per employee for courses and certifications'
    )
    
    doc.add_heading('Data Protection & Compliance Framework', 2)
    
    doc.add_heading('Data Handling Procedures', 3)
    doc.add_paragraph(
        'Customer data classification: Confidential and strictly protected\n'
        'Access controls: Role-based access to production systems\n'
        'Encryption: AES-256 for data at rest, TLS 1.3 for data in transit\n'
        'Backup procedures: Daily encrypted backups stored geographically dispersed\n'
        'Incident response: 24-hour incident response team with documented procedures\n'
        'Third-party vendors: Vetted and contractually bound to same data protection standards'
    )
    
    doc.add_heading('Security Measures', 3)
    doc.add_paragraph(
        'Access Authorization: Multi-factor authentication required for all systems\n'
        'Audit Logging: Complete audit trail of all data access with 1-year retention\n'
        'Penetration Testing: Annual security audits by external certified firms\n'
        'Employee Training: Mandatory yearly cybersecurity awareness training\n'
        'Breach Protocol: Incident reporting within 24 hours, government notification if required'
    )
    
    doc.add_heading('Employee Conduct Standards', 3)
    doc.add_paragraph(
        '1. All employees must complete data protection training before system access\n'
        '2. Confidentiality agreements signed upon employment with annual renewal\n'
        '3. Strict prohibition on sharing access credentials or passwords\n'
        '4. No personal use of company systems for external work or competitive advantage\n'
        '5. Violations subject to disciplinary action up to and including termination'
    )
    
    doc.add_heading('Compliance Responsibilities', 2)
    doc.add_paragraph(
        'Data Protection Officer: Designated and has direct board access for reporting\n'
        'Compliance Team: 5 personnel dedicated to regulatory adherence and audits\n'
        'Internal Audits: Quarterly reviews of data handling and security procedures\n'
        'External Compliance: Annual certification audits and regulatory filings\n'
        'Documentation: Comprehensive records of all policies, procedures, and training'
    )
    
    doc.add_heading('Business Continuity', 2)
    doc.add_paragraph(
        'Disaster Recovery Plan: Tested semi-annually with RTO of 4 hours, RPO of 1 hour\n'
        'Alternative Office Location: Secondary location with connectivity tested monthly\n'
        'Key Personnel Redundancy: Critical functions have documented succession plans\n'
        'Customer Communication: Clear protocols for notifying affected parties of any disruptions'
    )
    
    doc.add_paragraph(
        '\n--- End of Document ---\n'
        'Document prepared: 2024\n'
        'Last reviewed: April 2024\n'
        'Next review date: October 2024'
    )
    
    return doc


def create_document_2_manufacturing():
    """
    Document 2: Manufacturing Company Scenario
    Company: ABC Manufacturing Norway AS
    Topics: Workplace safety, worker schedules, equipment standards
    Related to: Workplace Act & Safety Regulations
    """
    doc = Document()
    
    title = doc.add_heading('Manufacturing Operations & Safety Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Company Information', 1)
    doc.add_paragraph('Company Name: ABC Manufacturing Norway AS')
    doc.add_paragraph('Establishment Year: 2008')
    doc.add_paragraph('Number of Employees: 143')
    doc.add_paragraph('Industry: Industrial Component Manufacturing')
    doc.add_paragraph('Main Operations: Precision machining, metal fabrication, assembly')
    doc.add_paragraph('Primary Location: Bergen, Norway')
    
    doc.add_heading('Production Overview', 2)
    doc.add_paragraph(
        'ABC Manufacturing Norway specializes in manufacturing precision metal components for automotive, '
        'marine, and industrial sectors. The company operates a modern 12,000 sqm facility with automated machinery, '
        'manual workstations, and quality control laboratories. The manufacturing process involves cutting, forming, '
        'heat treatment, machining, and final assembly of components for international clients.'
    )
    
    doc.add_heading('Workforce Structure', 2)
    
    doc.add_heading('Employee Categories & Roles', 3)
    doc.add_paragraph(
        'Production/Shop Floor Workers: 89 employees\n'
        '- Machine operators (CNC machines, manually operated equipment)\n'
        '- Assembly workers performing hand and power tool operations\n'
        '- Quality inspectors and equipment maintenance technicians\n'
        '\n'
        'Administrative & Planning: 32 employees\n'
        '- Production planning and scheduling\n'
        '- Quality assurance and documentation\n'
        '- Procurement and supply chain management\n'
        '- Sales and customer liaison\n'
        '\n'
        'Management & Supervision: 22 employees\n'
        '- Production floor supervisors (3 shifts)\n'
        '- Department managers\n'
        '- Safety coordinator (dedicated)\n'
        '- Executive management'
    )
    
    doc.add_heading('Work Schedules & Shifts', 3)
    doc.add_paragraph(
        'Shift System: Three rotating shifts (06:00-14:30, 14:00-22:30, 22:00-06:30)\n'
        'Days per cycle: 14-day rotation (7 days on, 7 days off)\n'
        'Break periods: 30-minute meal break, two 15-minute rest breaks per 8-hour shift\n'
        'Maximum weekly hours: 40 hours (including breaks)\n'
        'Overtime policy: Authorized by supervisor, compensated at 1.5x rate or time off\n'
        'Rest between shifts: Minimum 11 hours between end of one shift and start of next'
    )
    
    doc.add_heading('Workplace Safety Program', 2)
    
    doc.add_heading('Safety Hazard Identification', 3)
    doc.add_paragraph(
        'Machinery hazards:\n'
        '- Rotating equipment with pinch points\n'
        '- High-temperature processes (heat treatment up to 1200°C)\n'
        '- Hydraulic systems with stored energy\n'
        '- Compressed air systems (8 bar working pressure)\n'
        '\n'
        'Environmental hazards:\n'
        '- Cutting oil mist and metal dust generation\n'
        '- Noise levels (85-95 dB in production areas)\n'
        '- Chemical exposure (coolants, cleaning solvents)\n'
        '- Heat stress in forging and heat treatment sections'
    )
    
    doc.add_heading('Personal Protective Equipment (PPE)', 3)
    doc.add_paragraph(
        'Required on shop floor (all workers):\n'
        '- Safety helmets (hard hats)\n'
        '- Steel-toed safety boots (oil-resistant)\n'
        '- Safety glasses with side protection\n'
        '- High-visibility vests\n'
        '\n'
        'Department-specific PPE:\n'
        '- Heat treatment area: Fire-resistant aprons, heat-resistant gloves, face shield\n'
        '- Machining area: Hearing protection (earplugs, earmuffs), work gloves\n'
        '- Chemical handling: Nitrile gloves, chemical-resistant apron, respirator (when required)\n'
        '- Assembly area: Work gloves, apron, wrist protection'
    )
    
    doc.add_heading('Equipment Safety Standards', 3)
    doc.add_paragraph(
        'Machine inspection schedule:\n'
        '- Daily pre-shift inspections by operators\n'
        '- Weekly detailed inspections by maintenance team\n'
        '- Monthly comprehensive safety audits by certified technician\n'
        '- Annual third-party certification for high-risk equipment (heat treat furnaces)\n'
        '\n'
        'Safety interlocks & guards:\n'
        '- All CNC machines equipped with emergency stops within reach\n'
        '- Protective guards on all rotating equipment (verified monthly)\n'
        '- Hydraulic systems pressure relief set at rated capacity\n'
        '- Control panels locked when machinery in operation\n'
        '- Isolation procedures documented and posted at each workstation'
    )
    
    doc.add_heading('Safety Training & Competence', 3)
    doc.add_paragraph(
        'New employee onboarding: 3-day workplace safety induction (mandatory)\n'
        'Annual safety refresher: 4 hours per year for all shop floor workers\n'
        'Machine-specific training: Certification required before operating each equipment type\n'
        'Chemical handling: COSHH assessment training for workers in chemical exposure areas\n'
        'Supervisor training: Additional 16 hours annually on hazard recognition and incident response\n'
        'First aid: 15 employees trained and certified (refreshed every 2 years)'
    )
    
    doc.add_heading('Incident Reporting & Management', 3)
    doc.add_paragraph(
        'Reporting requirement: All near-misses and incidents reported to supervisor within 24 hours\n'
        'Investigation: Formal investigation for any incident with injury or damage within 48 hours\n'
        'Root cause analysis: Required for incidents causing lost work time\n'
        'Implementation: Corrective actions documented and tracked to completion\n'
        'Communication: Safety briefings on significant incidents at shift start meetings\n'
        'Recordkeeping: Complete incident database maintained for 7 years'
    )
    
    doc.add_heading('Health & Hygiene', 3)
    doc.add_paragraph(
        'Sanitation facilities: Adequate wash stations with hot water and soap\n'
        'Changing facilities: Separate areas for work clothes and personal items\n'
        'Occupational health: Annual health screening for workers in hazardous areas\n'
        'Hazard communication: Safety Data Sheets (SDS) available in Norwegian for all chemicals\n'
        'Hearing conservation: Baseline and annual audiometric testing for exposed workers'
    )
    
    doc.add_heading('Workplace Rights & Worker Protection', 2)
    doc.add_paragraph(
        'Right to stop work: Workers can pause operations if safety issue discovered\n'
        'Worker representation: Safety committee meets monthly with worker representatives\n'
        'No retaliation: Protection against dismissal for raising safety concerns\n'
        'Medical support: All work-related injuries covered by company insurance\n'
        'Return to work: Gradual return-to-work program for injured employees'
    )
    
    doc.add_heading('Management Responsibility', 2)
    doc.add_paragraph(
        'Safety budget: Dedicated annual budget for equipment upgrades and training\n'
        'Safety officer accountability: Safety performance linked to manager evaluations\n'
        'Compliance audits: Quarterly internal audits and annual external safety inspections\n'
        'Reporting hierarchy: Safety issues escalated to management within 24 hours\n'
        'Continuous improvement: Annual review of safety performance with target setting'
    )
    
    doc.add_paragraph(
        '\n--- End of Document ---\n'
        'Document prepared: 2024\n'
        'Last reviewed: April 2024\n'
        'Next review date: October 2024'
    )
    
    return doc


def create_document_3_finance():
    """
    Document 3: Financial Services Company Scenario
    Company: Nordic Finance Corp
    Topics: Financial reporting, accounting standards, audit requirements
    Related to: Accounting & Financial Regulations
    """
    doc = Document()
    
    title = doc.add_heading('Financial Services Operations & Accounting Framework', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Company Information', 1)
    doc.add_paragraph('Company Name: Nordic Finance Corp')
    doc.add_paragraph('Establishment Year: 2010')
    doc.add_paragraph('Number of Employees: 64')
    doc.add_paragraph('Industry: Financial Advisory & Investment Services')
    doc.add_paragraph('Regulatory Status: Licensed under Norwegian Financial Supervisory Authority (Finanstilsynet)')
    doc.add_paragraph('Principal Office: Oslo, Norway')
    
    doc.add_heading('Business Operations', 2)
    doc.add_paragraph(
        'Nordic Finance Corp provides comprehensive financial advisory services, investment management, and corporate '
        'financial planning to institutional and high-net-worth individual clients. The company manages client portfolios, '
        'provides tax planning advice, and offers financial risk consulting. All operations are conducted under strict regulatory '
        'oversight and Norwegian financial services regulations.'
    )
    
    doc.add_heading('Organizational Structure', 2)
    
    doc.add_heading('Department & Personnel Composition', 3)
    doc.add_paragraph(
        'Financial Advisory Team: 28 investment advisors and senior advisors\n'
        '- Bachelor minimum in finance, business, or related field\n'
        '- Most hold industry certifications (CFA, CFP)\n'
        '- Average client portfolio: 15-20 accounts per advisor\n'
        '\n'
        'Operations & Compliance: 18 employees\n'
        '- Compliance officers (4 dedicated)\n'
        '- Operations managers\n'
        '- Administrative support\n'
        '\n'
        'Finance & Accounting: 12 staff\n'
        '- Chief Financial Officer\n'
        '- Senior accountants (2)\n'
        '- Bookkeepers and accounting clerks (3)\n'
        '\n'
        'Management & Executive: 6 senior members\n'
        '- Chief Executive Officer\n'
        '- Chief Compliance Officer\n'
        '- Chief Risk Officer\n'
        '- Senior Management'
    )
    
    doc.add_heading('Financial Reporting Requirements', 2)
    
    doc.add_heading('Accounting Standards & Framework', 3)
    doc.add_paragraph(
        'Accounting standard: Norwegian Accounting Act (Regnskapsloven) with IFRS adaptation\n'
        'Chart of accounts: Structured according to Norwegian accounting classifications\n'
        'Currency: Primary currency NOK (Norwegian Krone)\n'
        'Fiscal year: Calendar year (January 1 - December 31)\n'
        'Reporting entity: Consolidated financial statements required\n'
        'Internal controls: COSO framework integrated into all financial processes'
    )
    
    doc.add_heading('Monthly Financial Procedures', 3)
    doc.add_paragraph(
        '1. Transaction Recording: All transactions recorded within 2 business days in accounting system\n'
        '2. Bank reconciliation: Performed by day 5 of following month\n'
        '3. Account verification: All subsidiary ledgers reconciled to general ledger\n'
        '4. Client account statements: Generated and sent to clients by 10th of month\n'
        '5. Management reporting: Internal P&L and balance sheet prepared by 15th\n'
        '6. Compliance review: transactions reviewed for regulatory compliance'
    )
    
    doc.add_heading('Annual Financial Reporting', 3)
    doc.add_paragraph(
        'Preparation timeline: Financial year close completed within 45 days\n'
        'Components:\n'
        '- Balance sheet (assets, liabilities, equity)\n'
        '- Income statement (revenues, expenses, profit)\n'
        '- Cash flow statement (operations, investing, financing)\n'
        '- Notes to financial statements (detailed disclosures)\n'
        '- Management commentary (business review and forward-looking statements)\n'
        '\n'
        'Audit: Mandatory external audit by registered auditor within 90 days of year-end\n'
        'Filing: Submitted to Brønnøysund Registry within statutory deadline (1 month after AGM)\n'
        'Public availability: Filed reports are public documents available through registry'
    )
    
    doc.add_heading('Internal Controls & Audit Procedures', 2)
    
    doc.add_heading('Control Environment', 3)
    doc.add_paragraph(
        'Chief Risk Officer: Responsible for internal control system design and operation\n'
        'Risk committee: Senior management reviews risks quarterly\n'
        'Segregation of duties: Separation of transaction approval, recording, and reconciliation\n'
        'Authorization limits: Documented approval thresholds based on transaction size and type\n'
        'Exception handling: All deviations from standard procedures documented and reviewed'
    )
    
    doc.add_heading('Compliance Monitoring', 3)
    doc.add_paragraph(
        'Regulatory requirements monitoring: Continuous review of updated regulations\n'
        'Client suitability: Quarterly review of all client portfolios for appropriateness\n'
        'Anti-money laundering: Enhanced due diligence for high-risk clients\n'
        'Conflict of interest: Disclosure and management of all related-party transactions\n'
        'Complaints handling: Customer complaints logged and investigation process documented\n'
        'Reporting to authorities: Suspicious activity reports filed as required'
    )
    
    doc.add_heading('Internal Audit Function', 3)
    doc.add_paragraph(
        'Annual audit plan: Developed by audit committee based on risk assessment\n'
        'Audit scope: Covers all material financial and operational processes\n'
        'Audit frequency:\n'
        '- High-risk areas: Quarterly reviews\n'
        '- Medium-risk areas: Semi-annual reviews\n'
        '- Low-risk areas: Annual reviews\n'
        '\n'
        'Report generation: Audit reports prepared with findings and recommendations\n'
        'Management response: Corrective actions documented with implementation timelines\n'
        'Follow-up: Verification that recommended actions completed'
    )
    
    doc.add_heading('Record Retention & Documentation', 2)
    doc.add_paragraph(
        'Financial records: Retained for minimum 5 years\n'
        'Client communications: All correspondence archived for 7 years (regulatory requirement)\n'
        'Audit documentation: Working papers retained for 3 years minimum\n'
        'Transaction records: Complete transaction history maintained with audit trail\n'
        'System backups: Daily backups stored securely with geographic redundancy\n'
        'Privacy compliance: All records stored securely with access logging'
    )
    
    doc.add_heading('External Audit Coordination', 2)
    doc.add_paragraph(
        'Auditor selection: Conducted through formal tender process annually\n'
        'Audit planning meeting: Held in December to discuss scope and timeline\n'
        'Information access: Complete access provided to auditors for testing\n'
        'Management letter: Auditor recommendations reviewed and implemented\n'
        'Reporting to board: Annual audit results presented to board/audit committee\n'
        'Fee transparency: Audit fees clearly disclosed in financial statements'
    )
    
    doc.add_heading('Reporting to Regulatory Authorities', 2)
    doc.add_paragraph(
        'Annual regulatory filing: Submitted to Finanstilsynet by deadline\n'
        'Quarterly statistics: Business data reported quarterly to regulatory body\n'
        'Capital adequacy: Minimum capital requirements verified and reported\n'
        'Personnel changes: Changes in key management reported to authorities\n'
        'Material events: Material business changes notified within statutory timeframe'
    )
    
    doc.add_paragraph(
        '\n--- End of Document ---\n'
        'Document prepared: 2024\n'
        'Last reviewed: April 2024\n'
        'Next review date: October 2024'
    )
    
    return doc


def create_document_4_realestate():
    """
    Document 4: Real Estate & Property Management Scenario
    Company: Property Management AS
    Topics: Tenant rights, property maintenance, lease agreements
    Related to: Real Estate & Property Law
    """
    doc = Document()
    
    title = doc.add_heading('Property Management Operations & Tenant Services', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Company Information', 1)
    doc.add_paragraph('Company Name: Property Management AS')
    doc.add_paragraph('Establishment Year: 2005')
    doc.add_paragraph('Number of Employees: 28')
    doc.add_paragraph('Industry: Commercial & Residential Property Management')
    doc.add_paragraph('Portfolio: 52 residential properties, 18 commercial units')
    doc.add_paragraph('Total Units Managed: 247 residential apartments, 42 commercial spaces')
    doc.add_paragraph('Principal Office: Stavanger, Norway')
    
    doc.add_heading('Property Portfolio Overview', 2)
    doc.add_paragraph(
        'Property Management AS manages a diverse portfolio of residential and commercial properties across several '
        'Norwegian municipalities. The company handles tenant relations, maintenance coordination, financial management, '
        'and regulatory compliance for property owners. The portfolio includes modern residential apartments, commercial offices, '
        'and mixed-use buildings.'
    )
    
    doc.add_heading('Tenant Rights & Lease Management', 2)
    
    doc.add_heading('Standard Lease Terms', 3)
    doc.add_paragraph(
        'Lease agreements: Compliant with Norwegian Tenancies Act (Husleieloven)\n'
        'Lease duration: Typically 1-3 year terms with automatic renewal unless terminated\n'
        'Notice period: 3 months notice required by either party for termination\n'
        'Deposit: Security deposit equal to 3 months rent (held in separate account)\n'
        'Rent increase: Maximum increase per year capped at inflation + 2% or as negotiated\n'
        'Rent payment: Due on first day of each month (automatic transfer preferred)'
    )
    
    doc.add_heading('Tenant Obligations & Rights', 3)
    doc.add_paragraph(
        'TENANT OBLIGATIONS:\n'
        '- Timely rent payment without deductions\n'
        '- Maintain property in good condition (normal wear and tear excepted)\n'
        '- Report maintenance issues within reasonable timeframe\n'
        '- Permit access for inspections (with 24 hours notice) and emergency repairs\n'
        '- Not alter property structure or decoration without written permission\n'
        '- Not sublet without property owner consent\n\n'
        'TENANT RIGHTS:\n'
        '- Right to occupy property for agreed duration (cannot be evicted without cause)\n'
        '- Right to peaceful enjoyment without interference\n'
        '- Right to return of security deposit (less documented damages) upon departure\n'
        '- Right to reasonable notice for any lease modifications\n'
        '- Right to access complaint and dispute resolution procedures'
    )
    
    doc.add_heading('Property Maintenance Program', 2)
    
    doc.add_heading('Maintenance Responsibilities & Schedule', 3)
    doc.add_paragraph(
        'PROPERTY MANAGEMENT RESPONSIBILITIES:\n'
        '- Roof, external walls, and foundation (major structural elements)\n'
        '- Building systems (electrical, plumbing, HVAC, fire safety)\n'
        '- Common areas (corridors, stairways, parking, gardens)\n'
        '- Exterior maintenance (grounds, drainage, landscaping)\n\n'
        'TENANT RESPONSIBILITIES:\n'
        '- Interior maintenance (walls, flooring, fixtures in unit)\n'
        '- Cleaning and hygiene\n'
        '- Minor repairs (light bulbs, door handles, etc.)\n'
        '- Fixture maintenance (sinks, faucets, locks)'
    )
    
    doc.add_heading('Preventive Maintenance Schedule', 3)
    doc.add_paragraph(
        'Weekly: Common area inspection and cleaning\n'
        'Monthly: Building systems check (HVAC filters, drainage)\n'
        'Quarterly: Structural inspection (roof, external walls)\n'
        'Semi-annual: Electrical and plumbing systems review\n'
        'Annual: Comprehensive safety inspection, fire equipment testing\n'
        'Biennial: Professional HVAC servicing, window/door sealing\n'
        'Triennial: Facade inspection, structural assessment'
    )
    
    doc.add_heading('Emergency Response Procedures', 3)
    doc.add_paragraph(
        'On-call emergency service: Available 24/7 for critical issues\n'
        'Response time: Same-day response for water leaks, electrical, heating failures\n'
        '48-hour response: Other maintenance issues\n'
        'Tenant reporting: Emergency hotline, email, and online portal available\n'
        'Contractor network: Pre-vetted licensed contractors on speed dial\n'
        'Documentation: All repairs logged with photos and invoices retained'
    )
    
    doc.add_heading('Financial Management', 2)
    
    doc.add_heading('Rent Collection & Accounting', 3)
    doc.add_paragraph(
        'Rent collection: Electronic bank transfers confirmed by 5th of month\n'
        'Late payment policy: Grace period 5 days, then accumulating late fees\n'
        '- Late fee: 1% of monthly rent per week overdue (max 10% of rent)\n'
        '- Legal action: Initiated if payment 30 days overdue\n\n'
        'Accounting procedures:\n'
        '- Monthly reconciliation of all accounts\n'
        '- Separate holding accounts for deposits\n'
        '- Interest on deposits accrued quarterly\n'
        '- Annual financial statements distributed to property owners'
    )
    
    doc.add_heading('Operating Expenses & Maintenance Budget', 3)
    doc.add_paragraph(
        'Expense categories:\n'
        '- Property tax and insurance (owner responsibility)\n'
        '- Maintenance and repairs (allocated to tenants proportionally)\n'
        '- Utilities (common areas only - tenant metered)\n'
        '- Management fees (standard industry rate ~8-12% of annual rent)\n\n'
        'Budget process:\n'
        '- Annual budget developed by September for following year\n'
        '- Approved by property owner board\n'
        '- Presented to tenants for information (advance notice 6 weeks)\n'
        '- Monthly reporting of actual vs. budgeted expenses'
    )
    
    doc.add_heading('Regulatory Compliance & Documentation', 2)
    
    doc.add_heading('Legal Compliance Requirements', 3)
    doc.add_paragraph(
        'Lease documentation: All leases properly executed and registered\n'
        'Tenancy register: Maintained with all tenant information and lease terms\n'
        'Dispute resolution: Formal process for rent disputes (mediation before court)\n'
        'Inspection records: Property inspections documented with photos\n'
        'Maintenance documentation: All work orders, estimates, and invoices filed\n'
        'Tenant confidentiality: Personal information protected per data protection regulations'
    )
    
    doc.add_heading('Insurance & Risk Management', 3)
    doc.add_paragraph(
        'Building insurance: Comprehensive property insurance covering structure and permanent fixtures\n'
        'Liability insurance: Coverage for injuries on property (min 10 million NOK)\n'
        'Tenant requirements: Some properties require tenant content insurance\n'
        'Professional indemnity: Management company insurance for professional liability\n'
        'Risk assessment: Annual review of insurance adequacy'
    )
    
    doc.add_heading('Health & Safety Standards', 3)
    doc.add_paragraph(
        'Fire safety: Annual fire alarm testing, escape route inspection\n'
        'Electrical safety: Regular inspections by certified electrician (every 5 years)\n'
        'Gas systems: Annual safety certification if applicable\n'
        'Asbestos management: Survey and registered safely if present\n'
        'Hazardous materials: Proper handling and disposal following regulations\n'
        'Accessibility: Common areas comply with accessibility standards where applicable'
    )
    
    doc.add_heading('Dispute Resolution & Tenant Relations', 2)
    doc.add_paragraph(
        'Complaint handling: Tenant complaints acknowledged within 48 hours\n'
        'Investigation process: Fair and documented investigation of all complaints\n'
        'Resolution timeline: Resolution offered within 30 days where possible\n'
        'Escalation path: Disputes escalated to management company director if unresolved\n'
        'Mediation: Option for professional mediation available\n'
        'Record keeping: All complaints and resolutions documented for 5 years'
    )
    
    doc.add_paragraph(
        '\n--- End of Document ---\n'
        'Document prepared: 2024\n'
        'Last reviewed: April 2024\n'
        'Next review date: October 2024'
    )
    
    return doc


def main():
    """Generate all test documents in DOCX format."""
    
    # Create test documents directory if it doesn't exist
    test_docs_dir = Path(__file__).parent / "test_docs"
    test_docs_dir.mkdir(exist_ok=True)
    
    print("📄 Creating test documents in DOCX format...\n")
    
    # Document 1: Tech Company
    print("Creating Document 1: Tech Company Compliance...")
    doc1 = create_document_1_tech_company()
    doc1_path = test_docs_dir / "test_doc_1_tech_company.docx"
    doc1.save(str(doc1_path))
    print(f"   ✅ Saved: {doc1_path}")
    
    # Document 2: Manufacturing Company
    print("Creating Document 2: Manufacturing Company...")
    doc2 = create_document_2_manufacturing()
    doc2_path = test_docs_dir / "test_doc_2_manufacturing.docx"
    doc2.save(str(doc2_path))
    print(f"   ✅ Saved: {doc2_path}")
    
    # Document 3: Finance Company
    print("Creating Document 3: Finance Company...")
    doc3 = create_document_3_finance()
    doc3_path = test_docs_dir / "test_doc_3_finance.docx"
    doc3.save(str(doc3_path))
    print(f"   ✅ Saved: {doc3_path}")
    
    # Document 4: Real Estate Company
    print("Creating Document 4: Real Estate Company...")
    doc4 = create_document_4_realestate()
    doc4_path = test_docs_dir / "test_doc_4_realestate.docx"
    doc4.save(str(doc4_path))
    print(f"   ✅ Saved: {doc4_path}")
    
    print("\n✅ All 4 test documents created successfully!")
    print(f"\n📁 Location: {test_docs_dir}")
    print("\nDocument Summary:")
    print("1. test_doc_1_tech_company.docx  - Software company compliance scenario")
    print("2. test_doc_2_manufacturing.docx - Manufacturing company safety scenario")
    print("3. test_doc_3_finance.docx       - Finance company accounting scenario")
    print("4. test_doc_4_realestate.docx    - Property management scenario")


if __name__ == "__main__":
    main()
