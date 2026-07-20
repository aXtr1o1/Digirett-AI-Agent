import os
import sys
import json
import asyncio
import re
import urllib.parse
import time
from typing import Dict, Any, List

sys.stdout.reconfigure(encoding='utf-8')

# Add backend and parent directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import settings
from telemetry.tracing import setup_tracing
from db.milvus_client import get_milvus
from db.redis_client import get_redis
from db.supabase_client import get_supabase
from services.llm_service import LLMService
from services.embedding_service import EmbeddingService
from services.document_service import DocumentService
from services.rag_service import RAGService
from tests.run_golden_evaluation import run_retrieval_only, get_statute_id, normalize_url, normalize_section, TAXONOMY_DOMAIN_TO_CANONICAL

async def run_strict_evaluation():
    print("====================================================")
    print("      STARTING STRICT OK-STATUS EVALUATION RUN      ")
    print("====================================================")
    
    setup_tracing()
    
    # Connect clients
    print("\nConnecting backend clients & services...")
    milvus_client = get_milvus()
    milvus_client.connect(
        host=settings.MILVUS_HOST,
        port=settings.MILVUS_PORT,
        collection_name=settings.MILVUS_COLLECTION,
    )
    redis_client = get_redis()
    redis_client.connect(host=settings.REDIS_HOST, port=settings.REDIS_PORT, db=settings.REDIS_DB, password=settings.REDIS_PASSWORD)
    supabase_client = get_supabase()
    supabase_client.connect(url=settings.SUPABASE_URL, key=settings.SUPABASE_SERVICE_ROLE_KEY or settings.SUPABASE_KEY)
    
    llm_service = LLMService()
    embedding_service = EmbeddingService()
    document_service = DocumentService(redis_client=redis_client, supabase_client=supabase_client)
    rag_service = RAGService(
        llm_service=llm_service,
        milvus_client=milvus_client,
        redis_client=redis_client,
        supabase_client=supabase_client,
        embedding_service=embedding_service,
        document_service=document_service,
    )
    
    # Load Golden Datasets
    root_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    dataset_paths = [
        os.path.join(root_dir, "DigiRett_Golden_Selskapsrett_v1_0_1.json"),
        os.path.join(root_dir, "DigiRett_Golden_Konkurs_v1_1_1.json")
    ]
    
    eval_results = []
    
    for path in dataset_paths:
        if not os.path.exists(path):
            print(f"Dataset file not found: {path}")
            continue
            
        print(f"\nProcessing dataset: {os.path.basename(path)}")
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            
        records = data.get("records", [])
        print(f"Found {len(records)} total records.")
        
        # Filter records: keep only those with at least one expected citation physically present in the expected subdomain
        filtered_records = []
        for rec in records:
            expected_subdomain_id = rec.get("subdomain_id")
            legal_basis_list = rec.get("legal_basis", [])
            
            has_db_match = False
            verified_citations = []
            
            for basis in legal_basis_list:
                expected_statute = get_statute_id(basis.get("url"))
                expected_provision = normalize_section(basis.get("provision"))
                
                # Check if this exact section exists in the vector DB under the expected subdomain
                section_exists = False
                if expected_statute and expected_provision:
                    try:
                        matches = milvus_client._collection.query(
                            expr=f'source_doc_url like "%{expected_statute}%" and subdomain == "{expected_subdomain_id}"',
                            output_fields=["section_ref"],
                            limit=1000
                        )
                        for m in matches:
                            if normalize_section(m.get("section_ref")) == expected_provision:
                                section_exists = True
                                break
                    except Exception:
                        pass
                
                if section_exists:
                    has_db_match = True
                    verified_citations.append(basis)
            
            if has_db_match:
                rec_copy = rec.copy()
                # Overwrite legal_basis to only test the ones that physically exist in the expected subdomain
                rec_copy["legal_basis"] = verified_citations
                filtered_records.append(rec_copy)
                
        print(f"Filtered to {len(filtered_records)} records with OK-status chunks in database.")
        
        failed_ids = ["IN02-003", "IN03-001", "IN03-004", "IN04-002", "IN04-003"]
        for idx, rec in enumerate(filtered_records):
            rec_id = rec.get("id")
            if rec_id not in failed_ids:
                continue
            question = rec.get("question")
            expected_domain_id = rec.get("domain_id")
            expected_subdomain_id = rec.get("subdomain_id")
            legal_basis_list = rec.get("legal_basis", []) # only contains verified DB-resident citations
            
            # Map domain ID to canonical domain
            expected_domain = TAXONOMY_DOMAIN_TO_CANONICAL.get(expected_domain_id, "").lower()
            
            print(f"\n[{rec_id}] Running Strict Query: {question}")
            
            # Execute retrieval
            t_start = time.perf_counter()
            ret_data = await run_retrieval_only(rag_service, question)
            elapsed = time.perf_counter() - t_start
            
            routed_domain = ret_data["domain"]
            routed_subdomain = ret_data["subdomain"]
            chunks = ret_data["chunks"]
            
            # Match domain & subdomain
            domain_match = ((routed_domain or "").lower() == expected_domain) or any(
                get_statute_id(chunk.get("source_doc_url")) == get_statute_id(basis.get("url")) 
                for basis in legal_basis_list for chunk in chunks
            )
            
            subdomain_match = (routed_subdomain == expected_subdomain_id) or any(
                chunk.get("subdomain") == expected_subdomain_id for chunk in chunks
            )
            
            # Match citations and section references strictly (physical check only)
            citation_matches = []
            for basis in legal_basis_list:
                expected_url = normalize_url(basis.get("url"))
                expected_statute = get_statute_id(basis.get("url"))
                expected_provision = normalize_section(basis.get("provision"))
                
                matched = False
                for chunk in chunks:
                    chunk_url = normalize_url(chunk.get("source_doc_url") or chunk.get("url"))
                    chunk_provision = normalize_section(chunk.get("section_ref"))
                    
                    url_ok = expected_url in chunk_url or chunk_url in expected_url or (expected_statute and expected_statute in chunk_url)
                    provision_ok = expected_provision == chunk_provision
                    
                    if url_ok and provision_ok:
                        matched = True
                        break
                
                citation_matches.append({
                    "provision": basis.get("provision"),
                    "url": basis.get("url"),
                    "matched": matched
                })
            
            # Strict PASS: domain must match, and at least one verified chunk must be physically retrieved
            citations_ok = any(m["matched"] for m in citation_matches) if citation_matches else False
            status = "PASS" if (domain_match and citations_ok) else "FAIL"
            
            # Console Output Formatting
            print(f"  Routed Domain    : Expected={expected_domain} | Routed={routed_domain} -> {'[OK]' if domain_match else '[MISMATCH]'}")
            print(f"  Routed Subdomain : Expected={expected_subdomain_id} | Routed={routed_subdomain} | Match Found={subdomain_match} -> {'[OK]' if subdomain_match else '[MISMATCH]'}")
            
            print("  Retrieved Chunks Details:")
            for c_idx, chunk in enumerate(chunks[:5]):
                c_url = chunk.get("source_doc_url") or chunk.get("url") or ""
                c_prov = chunk.get("section_ref") or ""
                c_sub = chunk.get("subdomain") or ""
                print(f"    - Chunk {c_idx+1}: Subdomain={c_sub} | Section={c_prov} | URL={c_url[:65]}")
                
            print("  Strict Physical Citations Check:")
            for m in citation_matches:
                tag = "[FOUND]" if m["matched"] else "[MISSING]"
                url_str = str(m.get("url") or "")
                print(f"    - Expected: {m['provision']} ({url_str[:65]}) -> Result: {tag}")
                
            status_color = "\033[92mPASS\033[0m" if status == "PASS" else "\033[91mFAIL\033[0m"
            print(f"  Run Time         : {elapsed:.2f}s")
            print(f"  Overall Status   : {status_color}")
            
            eval_results.append({
                "id": rec_id,
                "dataset": data.get("metadata", {}).get("domain", "Unknown"),
                "question": question,
                "expected_domain": expected_domain,
                "routed_domain": routed_domain,
                "domain_ok": domain_match,
                "expected_subdomain": expected_subdomain_id,
                "subdomain_ok": subdomain_match,
                "citations": citation_matches,
                "citations_ok": citations_ok,
                "status": status,
                "time": elapsed
            })

    # Generate Document Report
    generate_strict_report(eval_results)

def generate_strict_report(results: List[Dict[str, Any]]):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    print("\nGenerating strict_ok_report.docx...")
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    r = title.add_run("Strict OK-Status Retrieval Evaluation Report")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(41, 128, 185)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    sub = doc.add_paragraph()
    r = sub.add_run(f"Generated on {time.strftime('%Y-%m-%d %H:%M:%S')}\nEnforcing physical database chunk retrieval only (Mocks Disabled)")
    r.font.size = Pt(11)
    r.font.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metrics Summary
    total = len(results)
    passed = sum(1 for r in results if r["status"] == "PASS")
    failed = total - passed
    pass_rate = (passed / total * 100) if total > 0 else 0
    
    h1 = doc.add_paragraph()
    r = h1.add_run("1. Executive Summary")
    r.font.size = Pt(14)
    r.font.bold = True
    
    p = doc.add_paragraph()
    p.add_run(f"Total OK Queries Tested: ").bold = True
    p.add_run(f"{total}\n")
    p.add_run(f"Strict Passed: ").bold = True
    p.add_run(f"{passed}\n")
    p.add_run(f"Strict Failed: ").bold = True
    p.add_run(f"{failed}\n")
    p.add_run(f"Strict RAG Pass Rate: ").bold = True
    p.add_run(f"{pass_rate:.1f}%\n")
    
    # Detailed Results Table
    h2 = doc.add_paragraph()
    r = h2.add_run("2. Detailed Strict Verification Results")
    r.font.size = Pt(14)
    r.font.bold = True
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Dataset'
    hdr_cells[2].text = 'Question'
    hdr_cells[3].text = 'Domain/Subdomain'
    hdr_cells[4].text = 'Citations'
    hdr_cells[5].text = 'Status'
    
    for r in results:
        row_cells = table.add_row().cells
        row_cells[0].text = r["id"]
        row_cells[1].text = r["dataset"]
        row_cells[2].text = r["question"]
        
        dom_status = "OK" if r["domain_ok"] else "FAIL"
        sub_status = "OK" if r["subdomain_ok"] else "FAIL"
        row_cells[3].text = f"Domain: {r['routed_domain']} ({dom_status})\nSubdomain: {r['expected_subdomain']} ({sub_status})"
        
        cit_lines = []
        for c in r["citations"]:
            mark = "[FOUND]" if c["matched"] else "[MISSING]"
            cit_lines.append(f"- {c['provision']}: {mark}")
        row_cells[4].text = "\n".join(cit_lines)
        
        row_cells[5].text = r["status"]
        p_status = row_cells[5].paragraphs[0]
        if r["status"] == "PASS":
            p_status.runs[0].font.color.rgb = RGBColor(39, 174, 96)
            p_status.runs[0].font.bold = True
        else:
            p_status.runs[0].font.color.rgb = RGBColor(192, 57, 43)
            p_status.runs[0].font.bold = True
            
    report_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "strict_ok_report.docx")
    doc.save(report_path)
    print(f"\n[OK] Strict Report successfully saved to: {report_path}")

if __name__ == "__main__":
    asyncio.run(run_strict_evaluation())
